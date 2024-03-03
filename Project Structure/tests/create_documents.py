from docx import Document

# 1
doc1 = Document()
doc1.core_properties.author = "Alice Smith"
doc1.core_properties.category = "Story"
doc1.core_properties.comments = "This is a story about a curious cat named Mittens."
doc1.core_properties.content_status = "Draft"
doc1.core_properties.created = "2024-03-03T00:00:00"
doc1.core_properties.identifier = "12345678-1234-1234-1234-123456789012"
doc1.core_properties.keywords = "cat, story, adventure"
doc1.core_properties.language = "en-US"
doc1.core_properties.last_modified_by = "Alice Smith"
doc1.core_properties.last_printed = "2024-03-03T00:00:00"
doc1.core_properties.modified = "2024-03-03T00:00:00"
doc1.core_properties.revision = 1
doc1.core_properties.subject = "A Cat's Adventure"
doc1.core_properties.title = "fluttering butterfly outside the window."
doc1.core_properties.version = "1.0"
doc1.add_heading("Alice Smith's Cat", level=1)
paragraph1 = "In a cozy little house nestled amidst the bustling city, lived a ginger cat named Mittens, the beloved companion of Alice Smith. She was known for her playful antics and insatiable curiosity. One sunny afternoon, while basking in a patch of warm sunlight, Mittens spotted a fluttering butterfly outside the window. She leaped off the window sill, tail swishing excitedly, and darted towards the fluttering wings. Mittens pounced and batted at the butterfly with her paws, but the nimble creature fluttered away, leaving her wanting more. Undeterred, Mittens continued to explore the house, searching for new adventures."
doc1.add_paragraph(paragraph1)
doc1.save("Alice_Smith_Cat_Story.docx")

# 2
doc2 = Document()
doc2.core_properties.author = "Bob Johnson"
doc2.core_properties.category = "Animals"
doc2.core_properties.comments = "This is a story about a mischievous kitten."
doc2.core_properties.content_status = "Draft"
doc2.core_properties.created = "2024-03-03T00:00:00"
doc2.core_properties.identifier = "98765432-1234-1234-1234-123456789012"
doc2.core_properties.keywords = "cat, kitten, mischievous"
doc2.core_properties.language = "en-US"
doc2.core_properties.last_modified_by = "Bob Johnson"
doc2.core_properties.last_printed = "2024-03-03T00:00:00"
doc2.core_properties.modified = "2024-03-03T00:00:00"
doc2.core_properties.revision = 1
doc2.core_properties.subject = "A Kitten's Antics"
doc2.core_properties.title = "Bob Johnson's Kitten"
doc2.core_properties.version = "1.0"
doc2.add_heading("Bob Johnson's Kitten", level=1)
paragraph2 = "In a quaint farmhouse surrounded by rolling hills, lived a fluffy white kitten named Snowball, the playful companion of Bob Johnson. He was known for his insatiable curiosity and mischievous energy. One crisp autumn morning, while chasing a dust bunny across the living room floor, Snowball stumbled upon a ball of yarn left unattended. Entranced by the colorful strands, he pounced and batted at the yarn, unraveling it in a frenzy. He chased the loose yarn around the room, leaving a trail of tangled threads in his wake. Bob chuckled as he watched the kitten's playful chaos, knowing he would need to gather the yarn before it became another casualty of Snowball's adventures."
doc2.add_paragraph(paragraph2)
doc2.save("Bob_Johnson_Kitten_Story.docx")

# 3
doc3 = Document()
doc3.core_properties.author = "Charlie Brown"
doc3.core_properties.category = "Fiction"
doc3.core_properties.comments = "This is a story about a brave cat on a journey."
doc3.core_properties.content_status = "Draft"
doc3.core_properties.created = "2024-03-03T00:00:00"
doc3.core_properties.identifier = "11223344-1234-1234-1234-123456789012"
doc3.core_properties.keywords = "cat, journey, brave"
doc3.core_properties.language = "en-US"
doc3.core_properties.last_modified_by = "Charlie Brown"
doc3.core_properties.last_printed = "2024-03-03T00:00:00"
doc3.core_properties.modified = "2024-03-03T00:00:00"
doc3.core_properties.revision = 1
doc3.core_properties.subject = "The Cat's Quest"
doc3.core_properties.title = "Charlie Brown's Cat"
doc3.core_properties.version = "1.0"
doc3.add_heading("Charlie Brown's Cat", level=1)
paragraph3 = "In a charming little bookstore nestled down a quiet street, lived a sleek black cat named Midnight, the loyal companion of Charlie Brown. He was known for his quiet nature and adventurous spirit. One rainy afternoon, while exploring the city streets, Midnight stumbled upon an old map tucked away in a corner. The map depicted a mysterious island rumored to hold great treasures. Intrigued by the possibility of adventure, Midnight embarked on a perilous journey across stormy seas and treacherous terrain, facing danger and challenges at every turn. But with courage and determination, Midnight pressed on, guided by the promise of discovery and the hope of returning home a hero."
doc3.add_paragraph(paragraph3)
doc3.save("Charlie_Brown_Cat.docx")

# 4
doc4 = Document()
doc4.core_properties.author = "Jane Doe"
doc4.core_properties.category = "Story"
doc4.core_properties.comments = "This is a story about a pampered cat."
doc4.core_properties.content_status = "Draft"
doc4.core_properties.created = "2024-03-03T00:00:00"
doc4.core_properties.identifier = "55555555-1234-1234-1234-123456789012"
doc4.core_properties.keywords = "cat, pampered, story"
doc4.core_properties.language = "en-US"
doc4.core_properties.last_modified_by = "Jane Doe"
doc4.core_properties.last_printed = "2024-03-03T00:00:00"
doc4.core_properties.modified = "2024-03-03T00:00:00"
doc4.core_properties.revision = 1
doc4.core_properties.subject = "A Pampered Life"
doc4.core_properties.title = "Jane Doe's Cat"
doc4.core_properties.version = "1.0"
doc4.add_heading("Jane Doe's Cat", level=1)
paragraph4 = "In a luxurious apartment overlooking the bustling city, lived a Persian cat named Duchess, the prized possession of Jane Doe. She was known for her regal demeanor and sophisticated taste. Every morning, Duchess would be awakened by the gentle clinking of her silver food dish, filled with gourmet cat food. After a leisurely breakfast, she would spend her days lounging on plush cushions, basking in sunbeams, and occasionally batting at feather toys presented by her devoted owner. One afternoon, while perched on a window sill, Duchess observed a stray cat navigating the alleyway below. The stray, thin and scruffy, rummaged through a bin, searching for scraps. Duchess watched with a mixture of curiosity and disdain, her pampered life a stark contrast to the struggles of the alley cat."
doc4.add_paragraph(paragraph4)
doc4.save("Jane_Doe_Cat_Story.docx")

# 5
doc5 = Document()
doc5.core_properties.author = ""
doc5.core_properties.category = ""
doc5.core_properties.comments = ""
doc5.core_properties.content_status = ""
doc5.core_properties.created = ""  # ISO 8601 format
doc5.core_properties.identifier = ""
doc5.core_properties.keywords = ""
doc5.core_properties.language = ""
doc5.core_properties.last_modified_by = ""
doc5.core_properties.last_printed = ""  # ISO 8601 format
doc5.core_properties.modified = ""  # ISO 8601 format
doc5.core_properties.revision = 1
doc5.core_properties.subject = ""
doc5.core_properties.title = ""
doc5.core_properties.version = ""
doc5.save('test_document5.docx')
